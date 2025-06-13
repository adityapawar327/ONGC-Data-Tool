import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from io import BytesIO
from compare import read_file
from fuzzywuzzy import fuzz, process

def find_matching_column(df, target_fields, threshold=80):
    """Find matching column using fuzzy logic"""
    from fuzzywuzzy import process
    
    columns = list(df.columns)
    best_matches = []
    
    # If target_fields is a string, convert it to a list
    if isinstance(target_fields, str):
        target_fields = [target_fields]
        
    # Try each target field
    for field in target_fields:
        match = process.extractOne(field, columns)
        if match and match[1] >= threshold:
            best_matches.append((match[0], match[1]))
    
    # Return the best match if any were found
    if best_matches:
        return max(best_matches, key=lambda x: x[1])[0]
    return None

def group_data_by_hdd(df):
    """Group data by HDD No and create consolidated records"""
    grouped_data = []
    
    # Find required columns using fuzzy matching
    hdd_col = find_matching_column(df, ['HDD No', 'HDD', 'Media ID', 'Tape No'])
    area_col = find_matching_column(df, ['Area', 'Area Name', 'Block', 'Survey Area'])
    data_type_col = find_matching_column(df, ['Type of Data/Reports', 'Data Type', 'Type', 'Reports'])
    format_col = find_matching_column(df, ['File_Format', 'Format', 'File Format'])
    fsp_col = find_matching_column(df, ['FSP', 'First Shot Point', 'Start SP'])
    lsp_col = find_matching_column(df, ['LSP', 'Last Shot Point', 'End SP'])
    year_col = find_matching_column(df, ['Period of Data Generation', 'Year', 'Acquisition Year', 'Date'])
    
    if not hdd_col:
        st.error("Could not find HDD No column")
        return []
    
    # Group by HDD No
    hdd_groups = df.groupby(hdd_col)
    
    for hdd_no, group in hdd_groups:
        # Get the first occurrence for most fields
        first_row = group.iloc[0]
        
        # Create consolidated record
        record = {
            'HDD S/No': hdd_no,
            'Area Name / Block': first_row[area_col] if area_col else '',
            'Survey Type': '3D',  # Default or can be derived from other columns
            'Data': first_row[data_type_col] if data_type_col else '',
            'Format': first_row[format_col] if format_col else '',
            'Acquisition Year': first_row[year_col] if year_col else '',
            'Line Sequence': first_row[year_col] if year_col else '',  # Using year as requested
        }
        
        # Create SP field with FSP-LSP ranges
        sp_ranges = []
        if fsp_col and lsp_col:
            # Group consecutive rows and create ranges
            for idx, row in group.iterrows():
                fsp = row[fsp_col] if pd.notna(row[fsp_col]) else ''
                lsp = row[lsp_col] if pd.notna(row[lsp_col]) else ''
                if fsp and lsp:
                    sp_ranges.append(f"{fsp}-{lsp}")
        
        # If more than 7 ranges, show only first and last
        if len(sp_ranges) > 7:
            record['SP'] = f"{sp_ranges[0]}, {sp_ranges[-1]}"
        else:
            record['SP'] = ', '.join(sp_ranges) if sp_ranges else ''
        
        # Add record length and sample interval with fuzzy matching
        rec_length_col = find_matching_column(df, ['Rec. Length', 'Record Length', 'Recording Length', 'Length'])
        sample_interval_col = find_matching_column(df, ['Sample Interval', 'Sampling Rate', 'Sample Rate', 'Interval'])
        
        record['Rec. Length'] = first_row[rec_length_col] if rec_length_col else ''
        record['Sample Interval'] = first_row[sample_interval_col] if sample_interval_col else ''
        
        grouped_data.append(record)
    
    return grouped_data

def create_labeled_doc(grouped_records, logo_path):
    """Create a Word document with ONGC logo and structured field layout matching the exact format."""
    doc = Document()
    
    # Set up the document dimensions for landscape orientation
    section = doc.sections[0]
    section.page_width = Inches(16.54)  # A4 landscape width
    section.page_height = Inches(11.69)  # A4 landscape height
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    # Create cards - dynamic grid layout (4 cards per row for top section, 2 for bottom)
    total_cards = len(grouped_records)
    
    # First section: 4 cards per row
    cards_processed = 0
    while cards_processed < total_cards:
        # Determine cards per row based on remaining cards
        if total_cards - cards_processed >= 4:
            cards_per_row = 4
        elif total_cards - cards_processed == 3:
            cards_per_row = 3
        elif total_cards - cards_processed == 2:
            cards_per_row = 2
        else:
            cards_per_row = 1
        
        # Create table for this row
        table = doc.add_table(rows=1, cols=cards_per_row)
        table.autofit = False
        table.style = None
        
        # Set column widths based on cards per row - adjusted for proper 4-card fit
        if cards_per_row == 4:
            col_width = Inches(3.6)  # Slightly smaller to fit 4 cards
        elif cards_per_row == 3:
            col_width = Inches(4.8)
        elif cards_per_row == 2:
            col_width = Inches(7.2)
        else:
            col_width = Inches(14.4)
            
        for col in table.columns:
            col.width = col_width
        
        # Fill the row with cards
        for j in range(cards_per_row):
            if cards_processed + j < total_cards:
                record = grouped_records[cards_processed + j]
                cell = table.cell(0, j)
                cell._element.clear_content()
                
                # Add cell formatting
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                
                # Add cell margins for inner spacing
                tcMar = OxmlElement('w:tcMar')
                for margin in ['top', 'left', 'bottom', 'right']:
                    mar = OxmlElement(f'w:{margin}')
                    mar.set(qn('w:w'), '120')
                    mar.set(qn('w:type'), 'dxa')
                    tcMar.append(mar)
                tcPr.append(tcMar)
                
                # Add solid border to cell
                tcBdr = OxmlElement('w:tcBdr')
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '18')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), '000000')
                    tcBdr.append(border)
                tcPr.append(tcBdr)
                
                # Create header section with logo and title
                header_para = cell.add_paragraph()
                header_para.space_after = Pt(6)
                  # Add logo aligned to the left
                try:
                    logo_run = header_para.add_run()
                    logo_run.add_picture(logo_path, width=Inches(0.5))
                except Exception as e:
                    logo_run = header_para.add_run("ONGC")
                    logo_run.bold = True
                    logo_run.font.size = Pt(8)
                
                # Add space after logo
                header_para.add_run("  ")
                
                # Add SPIC, Mumbai in the same line
                title_run = header_para.add_run("SPIC, Mumbai")
                title_run.bold = True
                title_run.font.size = Pt(14)
                
                # Add Area subheading
                area_para = cell.add_paragraph()
                area_para.space_after = Pt(10)
                area_field_run = area_para.add_run('AREA')
                area_field_run.bold = True
                area_field_run.font.size = Pt(11)
                area_colon_run = area_para.add_run(' : ')
                area_colon_run.font.size = Pt(11)
                area_value_run = area_para.add_run(str(record.get('Area Name / Block', 'MB-OSN-2005-5 and 6')))
                area_value_run.font.size = Pt(11)
                
                # Add fields with consistent spacing and proper alignment
                field_order = ['Survey Type', 'Data', 'Format', 'Acquisition Year', 'SP', 'Rec. Length', 'Sample Interval', 'Line Sequence']
                
                for field in field_order:
                    para = cell.add_paragraph()
                    para.space_after = Pt(3)
                    
                    value = str(record.get(field, ''))
                    
                    # Format the field and value with exact spacing
                    field_run = para.add_run(field)
                    field_run.bold = True
                    field_run.font.size = Pt(10)
                    
                    colon_space_run = para.add_run(' : ')
                    colon_space_run.font.size = Pt(10)
                    
                    value_run = para.add_run(value)
                    value_run.font.size = Pt(10)
                
                # Add extra spacing before HDD S/No
                spacer_para = cell.add_paragraph()
                spacer_para.space_after = Pt(15)
                
                # Add HDD S/No at the bottom with center alignment
                hdd_para = cell.add_paragraph()
                hdd_run = hdd_para.add_run(f'HDD S/No: {record.get("HDD S/No", "")}')
                hdd_run.bold = True
                hdd_run.font.size = Pt(11)
                hdd_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        cards_processed += cards_per_row
        
        # Add spacing between rows
        if cards_processed < total_cards:
            spacer = doc.add_paragraph()
            spacer.space_after = Pt(20)
    
    return doc

def label_app():
    st.title("📑 Create Labeled Document - Grouped by HDD No")
    
    uploaded_files = st.file_uploader(
        "Upload Excel/CSV files",
        type=["xlsx", "xls", "csv"],
        accept_multiple_files=True,
        key="label_upload"
    )
    
    logo_path = "./public/ongc (1).jpg"
    
    if uploaded_files:
        all_grouped_records = []
        
        for file in uploaded_files:
            df = read_file(file)
            if df is not None:
                st.subheader(f"📄 Processing: {file.name}")
                
                # Show original data preview
                with st.expander(f"Original Data Preview: {file.name}"):
                    st.write("**Columns found:**")
                    st.write(list(df.columns))
                    st.write("**Sample original data:**")
                    st.dataframe(df.head(5))
                
                # Group data by HDD No
                grouped_records = group_data_by_hdd(df)
                
                if grouped_records:
                    all_grouped_records.extend(grouped_records)
                    
                    # Show grouped data preview
                    with st.expander(f"Grouped Data Preview: {file.name} ({len(grouped_records)} unique HDD records)"):
                        grouped_df = pd.DataFrame(grouped_records)
                        st.dataframe(grouped_df)
                else:
                    st.warning(f"No data could be grouped from {file.name}")
        
        if all_grouped_records:
            st.subheader(f"📋 Total Grouped Records: {len(all_grouped_records)}")
            
            # Show document preview
            st.subheader("🔍 Document Preview")
            
            # Display cards in dynamic grid layout
            total_cards = len(all_grouped_records)
            cards_processed = 0
            
            while cards_processed < total_cards:
                # Determine layout
                if total_cards - cards_processed >= 4:
                    cols_per_row = 4
                elif total_cards - cards_processed == 3:
                    cols_per_row = 3
                elif total_cards - cards_processed == 2:
                    cols_per_row = 2
                else:
                    cols_per_row = 1
                
                cols = st.columns(cols_per_row)
                
                for j in range(cols_per_row):
                    if cards_processed + j < total_cards:
                        record = all_grouped_records[cards_processed + j]
                        
                        with cols[j]:
                            # Create bordered preview card with exact styling
                            st.markdown("""
                                <div style="border: 2px solid black; padding: 12px; margin: 8px; background-color: white; font-family: Arial, sans-serif;">
                            """, unsafe_allow_html=True)
                            
                            # Header with logo and title
                            col1, col2 = st.columns([1, 3])
                            with col1:
                                try:
                                    st.image(logo_path, width=40)
                                except:
                                    st.markdown("**ONGC**")
                            with col2:
                                st.markdown("<h4 style='margin: 8px 0; padding-top: 8px; font-size: 16px; font-weight: bold;'>SPIC, Mumbai</h4>", unsafe_allow_html=True)
                            
                            # Area
                            area_name = record.get('Area Name / Block', 'MB-OSN-2005-5 and 6')
                            st.markdown(f"<p style='font-weight: bold; margin: 12px 0; font-size: 12px;'><strong>AREA</strong> : {area_name}</p>", unsafe_allow_html=True)
                            
                            # Display fields with exact formatting
                            field_order = ['Survey Type', 'Data', 'Format', 'Acquisition Year', 'SP', 'Rec. Length', 'Sample Interval', 'Line Sequence']
                            
                            for field in field_order:
                                value = record.get(field, '')
                                # Truncate long SP values for display
                                if field == 'SP' and len(str(value)) > 50:
                                    display_value = str(value)[:47] + "..."
                                else:
                                    display_value = str(value)
                                
                                st.markdown(f"<div style='margin: 4px 0; font-size: 11px; text-align: left;'><strong>{field}</strong> : {display_value}</div>", unsafe_allow_html=True)
                            
                            # Add spacing and HDD S/No at bottom center  
                            st.markdown("<div style='margin-top: 20px;'></div>", unsafe_allow_html=True)
                            
                            hdd_value = record.get('HDD S/No', '')
                            st.markdown(f"<div style='text-align: center; font-weight: bold; font-size: 12px;'>HDD S/No: {hdd_value}</div>", unsafe_allow_html=True)
                            
                            st.markdown("</div>", unsafe_allow_html=True)
                
                cards_processed += cols_per_row
            
            if st.button("🔖 Generate Labeled Document (Grouped Data)"):
                try:
                    doc = create_labeled_doc(all_grouped_records, logo_path)
                    
                    # Save to BytesIO
                    doc_io = BytesIO()
                    doc.save(doc_io)
                    doc_io.seek(0)
                    
                    st.success(f"✅ Document generated successfully with {len(all_grouped_records)} grouped records!")
                    
                    # Show summary statistics
                    st.info(f"""
                    **Summary:**
                    - Total unique HDD records: {len(all_grouped_records)}
                    - Cards will be arranged in dynamic grid layout
                    - Each card represents one grouped HDD record
                    """)
                    
                    # Offer download
                    st.download_button(
                        label="📥 Download Word Document (Grouped Data)",
                        data=doc_io,
                        file_name="labeled_document_grouped_data.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessing",
                    )
                    
                except Exception as e:
                    st.error(f"Error generating document: {e}")
                    st.write("Debug info:")
                    st.write(f"Number of grouped records: {len(all_grouped_records)}")
                    if all_grouped_records:
                        st.write("Sample record:")
                        st.json(all_grouped_records[0])
        else:
            st.warning("No valid grouped data found. Please check your files and try again.")
    else:
        st.info("Please upload Excel/CSV files to generate a labeled document with grouped data.")

if __name__ == "__main__":
    label_app()